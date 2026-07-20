#!/usr/bin/env python3
"""
Weekly password rotation for the CLC SAT Quiz site.

What this does, in order:
  1. Generates a new random shared password.
  2. Writes its SHA-256 hash into auth-config.js (the site reads this at runtime).
  3. Builds a small Word document with the site link + the new password,
     and encrypts it (password-protects it) with that same password.
  4. Uploads/replaces that Word document in a shared Google Drive folder.
  5. Emails the new password to the configured recipient(s).

Secrets/config are read from environment variables (see .github/workflows/rotate-password.yml):
  GMAIL_ADDRESS                 - sender Gmail address
  GMAIL_APP_PASSWORD            - 16-char Gmail App Password
  GDRIVE_SERVICE_ACCOUNT_JSON   - full JSON key of the Google service account
  GDRIVE_FOLDER_ID              - Drive folder ID to upload the Word doc into
  PASSWORD_RECIPIENT_EMAILS     - comma-separated list of emails to notify
"""
import hashlib
import io
import json
import os
import secrets
import smtplib
from email.mime.text import MIMEText

from docx import Document
from docx.shared import Pt, RGBColor
import msoffcrypto
from msoffcrypto.format.ooxml import OOXMLFile

from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload

SITE_URL = "https://stavros-tsal.github.io/clc-sat-quiz/"
AUTH_CONFIG_PATH = os.path.join(os.path.dirname(__file__), "..", "auth-config.js")
DOCX_FILENAME = "CLC SAT Quiz - Access.docx"

# Unambiguous alphabet: no 0/O, 1/l/I confusion.
PASSWORD_ALPHABET = "ABCDEFGHJKMNPQRSTUVWXYZabcdefghijkmnpqrstuvwxyz23456789"
PASSWORD_LENGTH = 10


def generate_password():
    return "".join(secrets.choice(PASSWORD_ALPHABET) for _ in range(PASSWORD_LENGTH))


def sha256_hex(text):
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def update_auth_config(password_hash):
    content = (
        "// Auto-updated weekly by GitHub Actions — do not edit by hand.\n"
        "// Holds the SHA-256 hash of the current shared student password.\n"
        f'const PASSWORD_HASH = "{password_hash}";\n'
    )
    with open(AUTH_CONFIG_PATH, "w", encoding="utf-8") as f:
        f.write(content)


def build_encrypted_docx(password):
    doc = Document()

    title = doc.add_paragraph()
    run = title.add_run("CLC SAT Quiz — Πρόσβαση")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x7B, 0x00, 0x00)

    doc.add_paragraph()

    p1 = doc.add_paragraph()
    p1.add_run("Site: ").bold = True
    p1.add_run(SITE_URL)

    p2 = doc.add_paragraph()
    p2.add_run("Password αυτής της εβδομάδας: ").bold = True
    run_pw = p2.add_run(password)
    run_pw.bold = True
    run_pw.font.size = Pt(14)

    doc.add_paragraph()
    note = doc.add_paragraph(
        "Το password αυτό ανανεώνεται αυτόματα κάθε Δευτέρα. "
        "Αυτό το αρχείο ενημερώνεται μαζί του."
    )
    note.runs[0].italic = True

    plain_buf = io.BytesIO()
    doc.save(plain_buf)
    plain_buf.seek(0)

    office_file = OOXMLFile(plain_buf)
    encrypted_buf = io.BytesIO()
    office_file.encrypt(password, encrypted_buf)
    encrypted_buf.seek(0)
    return encrypted_buf


def upload_to_drive(encrypted_buf, folder_id, service_account_json):
    info = json.loads(service_account_json)
    creds = service_account.Credentials.from_service_account_info(
        info, scopes=["https://www.googleapis.com/auth/drive"]
    )
    drive = build("drive", "v3", credentials=creds)

    media = MediaIoBaseUpload(
        encrypted_buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        resumable=False,
    )

    existing = drive.files().list(
        q=f"name='{DOCX_FILENAME}' and '{folder_id}' in parents and trashed=false",
        fields="files(id,name)",
        supportsAllDrives=True,
        includeItemsFromAllDrives=True,
    ).execute()
    files = existing.get("files", [])

    if not files:
        # Debug aid: list everything actually visible in that folder so a
        # naming/location mismatch is obvious in the Action logs instead of
        # failing silently on create().
        everything = drive.files().list(
            q=f"'{folder_id}' in parents and trashed=false",
            fields="files(id,name,mimeType)",
            supportsAllDrives=True,
            includeItemsFromAllDrives=True,
        ).execute()
        print(f"Looking for exact name: {DOCX_FILENAME!r}")
        print("Files visible to the service account in that folder:")
        for f in everything.get("files", []):
            print(f"  - {f['name']!r} (mimeType={f['mimeType']}, id={f['id']})")
        if not everything.get("files"):
            print("  (none — folder appears empty or not shared with the service account)")

    if files:
        file_id = files[0]["id"]
        drive.files().update(fileId=file_id, media_body=media, supportsAllDrives=True).execute()
        print(f"Updated existing Drive file: {file_id}")
    else:
        metadata = {"name": DOCX_FILENAME, "parents": [folder_id]}
        created = drive.files().create(
            body=metadata, media_body=media, fields="id", supportsAllDrives=True
        ).execute()
        print(f"Created new Drive file: {created.get('id')}")


def send_email(password, recipients, gmail_address, gmail_app_password):
    body = (
        f"Νέο password για το CLC SAT Quiz αυτής της εβδομάδας:\n\n"
        f"    {password}\n\n"
        f"Site: {SITE_URL}\n\n"
        f"Το ίδιο password ανοίγει και το Word doc στον Drive φάκελο "
        f"(\"{DOCX_FILENAME}\").\n\n"
        f"— Αυτόματο μήνυμα από το clc-sat-quiz repo"
    )
    msg = MIMEText(body, "plain", "utf-8")
    msg["Subject"] = "CLC SAT Quiz — Νέο password εβδομάδας"
    msg["From"] = gmail_address
    msg["To"] = ", ".join(recipients)

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
        smtp.login(gmail_address, gmail_app_password)
        smtp.sendmail(gmail_address, recipients, msg.as_string())


def main():
    gmail_address = os.environ["GMAIL_ADDRESS"]
    gmail_app_password = os.environ["GMAIL_APP_PASSWORD"]
    sa_json = os.environ["GDRIVE_SERVICE_ACCOUNT_JSON"]
    folder_id = os.environ["GDRIVE_FOLDER_ID"]
    recipients = [e.strip() for e in os.environ["PASSWORD_RECIPIENT_EMAILS"].split(",") if e.strip()]

    password = generate_password()
    password_hash = sha256_hex(password)
    print(f"Generated new password (hash only logged): {password_hash}")

    update_auth_config(password_hash)
    print("Updated auth-config.js")

    encrypted_docx = build_encrypted_docx(password)
    print("Built encrypted Word document")

    upload_to_drive(encrypted_docx, folder_id, sa_json)
    print("Uploaded to Google Drive")

    send_email(password, recipients, gmail_address, gmail_app_password)
    print(f"Emailed new password to: {', '.join(recipients)}")


if __name__ == "__main__":
    main()
